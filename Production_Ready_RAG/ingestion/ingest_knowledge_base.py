from pathlib import Path
import csv
import sys

# ============================================================
# Add project root to Python path
# ============================================================

PROJECT_ROOT = Path(
    __file__
).resolve().parent.parent

if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(
        0,
        str(PROJECT_ROOT)
    )


# ============================================================
# Project imports
# ============================================================

from models.document import Document
from embeddings.embedding_model import EmbeddingModel
from chunking.text_chunker import TextChunker
from vectordb.chroma_db import VectorDatabase

# ============================================================
# Configuration
# ============================================================

DATA_DIRECTORY = Path(
    "data/company_info"
)

VECTOR_DB_DIRECTORY = (
    "data/vector_db"
)

COLLECTION_NAME = (
    "hr_knowledge_base"
)


# ============================================================
# TXT Loader
# ============================================================

def load_txt(
    file_path: Path
) -> str:

    return file_path.read_text(
        encoding="utf-8"
    )


# ============================================================
# PDF Loader
# ============================================================

def load_pdf(
    file_path: Path
) -> str:

    from pypdf import PdfReader

    reader = PdfReader(
        str(file_path)
    )

    pages = []

    for page in reader.pages:

        text = page.extract_text()

        if text:

            pages.append(
                text
            )

    return "\n".join(pages)


# ============================================================
# DOCX Loader
# ============================================================

def load_docx(
    file_path: Path
) -> str:

    from docx import Document as DocxDocument

    doc = DocxDocument(
        str(file_path)
    )

    paragraphs = []

    for paragraph in doc.paragraphs:

        text = paragraph.text.strip()

        if text:

            paragraphs.append(
                text
            )

    return "\n".join(paragraphs)


# ============================================================
# CSV Loader
# ============================================================

def load_csv(
    file_path: Path
) -> str:

    rows = []

    with open(
        file_path,
        "r",
        encoding="utf-8-sig",
        newline=""
    ) as file:

        reader = csv.DictReader(
            file
        )

        for row in reader:

            fields = []

            for key, value in row.items():

                if value:

                    fields.append(
                        f"{key}: {value}"
                    )

            if fields:

                rows.append(
                    "\n".join(fields)
                )

    return "\n\n".join(rows)


# ============================================================
# File Loader
# ============================================================

def load_file(
    file_path: Path
) -> str:

    extension = (
        file_path.suffix.lower()
    )

    if extension == ".txt":

        return load_txt(
            file_path
        )

    if extension == ".pdf":

        return load_pdf(
            file_path
        )

    if extension == ".docx":

        return load_docx(
            file_path
        )

    if extension == ".csv":

        return load_csv(
            file_path
        )

    raise ValueError(
        f"Unsupported file type: {extension}"
    )


# ============================================================
# Discover and Load Documents
# ============================================================

def load_all_documents():

    documents = []

    supported_extensions = {
        ".txt",
        ".pdf",
        ".docx",
        ".csv"
    }

    files = sorted(
        DATA_DIRECTORY.rglob("*")
    )

    for file_path in files:

        if not file_path.is_file():
            continue

        if (
            file_path.suffix.lower()
            not in supported_extensions
        ):
            continue

        print(
            f"Loading: {file_path}"
        )

        try:

            content = load_file(
                file_path
            )

            if not content.strip():

                print(
                    f"Skipping empty file: "
                    f"{file_path}"
                )

                continue

            document = Document(
                content=content,
                metadata={
                    "source": str(file_path),
                    "file_name": file_path.name,
                    "file_type": (
                        file_path.suffix.lower()
                    )
                }
            )

            documents.append(
                document
            )

        except Exception as error:

            print(
                f"ERROR loading "
                f"{file_path}: {error}"
            )

    return documents


# ============================================================
# Main Ingestion Pipeline
# ============================================================

def main():

    print()
    print("=" * 70)
    print("HR KNOWLEDGE BASE INGESTION")
    print("=" * 70)

    print(
        f"\nData directory:"
        f" {DATA_DIRECTORY}"
    )

    # --------------------------------------------------------
    # 1. Load files
    # --------------------------------------------------------

    documents = load_all_documents()

    print()
    print(
        f"Documents loaded: "
        f"{len(documents)}"
    )

    if not documents:

        print(
            "No supported documents found."
        )

        return

    # --------------------------------------------------------
    # 2. Chunk documents
    # --------------------------------------------------------

    print()
    print(
        "Creating document chunks..."
    )

    chunker = TextChunker(
        chunk_size=1000,
        chunk_overlap=200
    )

    chunked_documents = (
        chunker.chunk_documents(
            documents
        )
    )

    print(
        f"Chunks created: "
        f"{len(chunked_documents)}"
    )

    # --------------------------------------------------------
    # 3. Load embedding model
    # --------------------------------------------------------

    print()
    print(
        "Loading embedding model..."
    )

    embedding_model = EmbeddingModel()

    # --------------------------------------------------------
    # 4. Generate embeddings
    # --------------------------------------------------------

    print()
    print(
        "Generating embeddings..."
    )

    embeddings = (
        embedding_model.embed_documents(
            chunked_documents
        )
    )

    print(
        f"Embeddings generated: "
        f"{len(embeddings)}"
    )

    # --------------------------------------------------------
    # 5. Store in ChromaDB
    # --------------------------------------------------------

    print()
    print(
        "Storing chunks in ChromaDB..."
    )

    vector_db = VectorDatabase(
        persist_directory=VECTOR_DB_DIRECTORY,
        collection_name=COLLECTION_NAME
    )

    vector_db.add_documents(
        documents=chunked_documents,
        embeddings=embeddings
    )

    # --------------------------------------------------------
    # 6. Summary
    # --------------------------------------------------------

    print()
    print("=" * 70)
    print("INGESTION COMPLETED SUCCESSFULLY")
    print("=" * 70)

    print(
        f"Documents : {len(documents)}"
    )

    print(
        f"Chunks    : {len(chunked_documents)}"
    )

    print(
        f"Collection: {COLLECTION_NAME}"
    )

    print(
        f"Vector DB : {VECTOR_DB_DIRECTORY}"
    )

    print("=" * 70)


if __name__ == "__main__":

    main()