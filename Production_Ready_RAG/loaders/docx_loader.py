from pathlib import Path

from docx import Document as DocxDocument

from models.document import Document


class DOCXLoader:
    """
    Loader for Microsoft Word (.docx) files.
    """

    @staticmethod
    def load(file_path: str) -> Document:
        """
        Reads a DOCX file and returns a Document object.
        """

        path = Path(file_path)

        doc = DocxDocument(path)

        paragraphs = []

        for paragraph in doc.paragraphs:

            text = paragraph.text.strip()

            if text:
                paragraphs.append(text)

        content = "\n".join(paragraphs)

        metadata = {
            "source": str(path),
            "file_name": path.name,
            "file_type": "docx"
        }

        return Document(
            content=content,
            metadata=metadata
        )